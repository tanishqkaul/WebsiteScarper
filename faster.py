import time
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from selenium import webdriver
from selenium.webdriver.chrome.options import Options

# --- CONFIGURATION ---
BASE_URL = "https://romakksilicones.com/"
WORD_FILENAME = "website_ROMAKK_optimized_selenium.docx"
MAX_WORKERS = 5 # Reduced workers slightly as browsers are more resource-intensive
WAIT_TIME = 2 # Increased wait time for dynamic content to load

# --- HELPER FUNCTIONS for DOCX (Unchanged) ---
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(11)

def add_table(doc, bs_table):
    rows = bs_table.find_all("tr")
    if not rows:
        return
    max_cols = 0
    for row in rows:
        max_cols = max(max_cols, len(row.find_all(['td', 'th'])))
    if max_cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                table.cell(i, j).text = cell.get_text(strip=True)

def add_list(doc, items, is_ordered):
    for i, item in enumerate(items, start=1):
        text = item.get_text(strip=True)
        style = 'List Number' if is_ordered else 'List Bullet'
        doc.add_paragraph(text, style=style)

# --- CORE LOGIC ---

def get_driver():
    """Initializes a headless Chrome WebDriver with optimizations."""
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1920,1080")
    # --- OPTIMIZATION: Disable image loading ---
    chrome_options.add_experimental_option("prefs", {"profile.managed_default_content_settings.images": 2})
    return webdriver.Chrome(options=chrome_options)

def parse_content_and_links(html_content, url):
    """Parses HTML to extract content structure and all internal links."""
    if not html_content:
        return [], set()

    soup = BeautifulSoup(html_content, "html.parser")
    
    links = set()
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith(('mailto:', 'tel:', 'javascript:')) or 'cdn-cgi' in href:
            continue
        full_url = urljoin(url, href)
        parsed = urlparse(full_url)
        if parsed.netloc == urlparse(BASE_URL).netloc:
            clean_url = parsed.scheme + "://" + parsed.netloc + parsed.path
            links.add(clean_url)

    content_structure = []
    for nav in soup.find_all('nav'):
        nav.decompose()
        
    processed_tags = []
    body_content = soup.body
    if body_content:
        for tag in body_content.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol'], recursive=True):
            if any(p in tag.parents for p in processed_tags):
                continue

            if tag.name.startswith('h'):
                try:
                    lvl = int(tag.name[1])
                    content_structure.append(('heading', tag.get_text(strip=True), min(lvl, 4)))
                except (ValueError, IndexError):
                    pass 
            elif tag.name == 'table':
                content_structure.append(('table', tag))
            elif tag.name == 'ul':
                content_structure.append(('list', tag.find_all('li', recursive=False), False))
            elif tag.name == 'ol':
                content_structure.append(('list', tag.find_all('li', recursive=False), True))
            elif tag.name == 'p' and tag.get_text(strip=True):
                content_structure.append(('paragraph', tag.get_text(strip=True)))

            processed_tags.append(tag)
            
    return content_structure, links

def worker(url):
    """
    Worker function for each thread. 
    Uses a dedicated Selenium driver to fetch and parse a page.
    """
    driver = get_driver()
    html_content = None
    try:
        driver.get(url)
        time.sleep(WAIT_TIME)  # Wait for dynamic content to load
        html_content = driver.page_source
    except Exception as e:
        print(f"Error fetching {url} with Selenium: {e}")
    finally:
        driver.quit() # Ensure the browser is closed to free resources
    
    return url, parse_content_and_links(html_content, url)

def main():
    """Main function to manage the crawling process."""
    doc = docx.Document()
    visited_urls = set()
    visited_lock = threading.Lock()
    
    start_time = time.time()

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        with visited_lock:
            visited_urls.add(BASE_URL)
        
        futures = {executor.submit(worker, BASE_URL)}
        
        while futures:
            for future in as_completed(futures):
                url, (content, new_links) = future.result()
                futures.remove(future)

                print(f"Processed: {url} ({len(visited_urls)} pages discovered)")

                # Write content to docx
                add_heading(doc, f"Page: {url}", level=1)
                for item_type, *data in content:
                    if item_type == 'heading':
                        add_heading(doc, data[0], level=data[1])
                    elif item_type == 'paragraph':
                        add_paragraph(doc, data[0])
                    elif item_type == 'table':
                        add_table(doc, data[0])
                    elif item_type == 'list':
                        add_list(doc, data[0], is_ordered=data[1])
                doc.add_page_break()

                for link in new_links:
                    with visited_lock:
                        if link not in visited_urls:
                            visited_urls.add(link)
                            futures.add(executor.submit(worker, link))
    
    doc.save(WORD_FILENAME)
    end_time = time.time()

    print("\n--- Crawling Complete ---")
    print(f"Saved content from {len(visited_urls)} pages to {WORD_FILENAME}")
    print(f"Total time taken: {end_time - start_time:.2f} seconds")

if __name__ == "__main__":
    main()
