import os
import time
from urllib.parse import urlparse, urljoin
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt

# Configuration
BASE_URL = "https://romakksilicones.com/"  # Change as needed
WORD_FILENAME = "website_romakk.docx"
WAIT_TIME = 0  # seconds to wait for JS/dynamic content

# Selenium Setup (Headless Chrome)
chrome_options = Options()
chrome_options.add_argument("--headless")
chrome_options.add_argument("--disable-gpu")
chrome_options.add_argument("--window-size=1920,1080")
driver = webdriver.Chrome(options=chrome_options)

visited = set()
doc = docx.Document()
header_footer_saved = False  # Flag to ensure header/footer saved once

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text):
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(11)

# --- MODIFIED FUNCTION ---
# This function is updated to handle irregular HTML tables, which was the source of the error.
def add_table(bs_table):
    """
    Adds a BeautifulSoup table to the Word document, handling tables with
    inconsistent numbers of columns across rows.
    """
    rows = bs_table.find_all("tr")
    if not rows:
        return

    # First, determine the maximum number of columns in any row.
    # This prevents the "index out of range" error for tables with irregular structures.
    max_cols = 0
    for row in rows:
        num_cells = len(row.find_all(['td', 'th']))
        if num_cells > max_cols:
            max_cols = num_cells

    # If the table has no columns, don't try to create it.
    if max_cols == 0:
        return

    # Create the table in the Word document with the maximum column count.
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid' # Optional: apply a default style for visibility

    # Populate the table, ensuring not to go out of bounds.
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            # This check is now safe because the table is created with max_cols.
            if j < max_cols:
                table.cell(i, j).text = cell.get_text(strip=True)
# --- END OF MODIFICATION ---

def add_list(items, is_ordered):
    for i, item in enumerate(items, start=1):
        text = item.get_text(strip=True)
        if is_ordered:
            doc.add_paragraph(f"{i}. {text}", style='List Number')
        else:
            doc.add_paragraph(text, style='List Bullet')

def parse_and_save_content(soup):
    global header_footer_saved
    # Save header and footer once, excluding navbar links from header content
    if not header_footer_saved:
        header = soup.find('header')
        if header:
            # Remove navbar from header before saving header content, if present
            navbar = header.find('nav')
            if navbar:
                navbar.decompose()
            add_heading('Header', level=1)
            for text in header.stripped_strings:
                add_paragraph(text)
        footer = soup.find('footer')
        if footer:
            add_heading('Footer', level=1)
            for text in footer.stripped_strings:
                add_paragraph(text)
        header_footer_saved = True

    # Remove navbar(s) from body to prevent duplicate link processing in content
    for nav in soup.find_all('nav'):
        nav.decompose()

    # Extract content in order respecting headings, tables, lists, paragraphs
    # We need to prevent double-processing of content inside tables/lists
    processed_tags = []
    for tag in soup.body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table', 'ul', 'ol'], recursive=True):
        # Skip if this tag is inside another tag we've already processed (e.g., a p inside a table cell)
        if any(p in tag.parents for p in processed_tags):
            continue

        if tag.name.startswith('h'):
            lvl = int(tag.name[1]) if tag.name[1].isdigit() else 1
            add_heading(tag.get_text(strip=True), min(lvl, 4))
        elif tag.name == 'table':
            add_table(tag)
        elif tag.name == 'ul':
            add_list(tag.find_all('li', recursive=False), is_ordered=False)
        elif tag.name == 'ol':
            add_list(tag.find_all('li', recursive=False), is_ordered=True)
        elif tag.name == 'p' and tag.get_text(strip=True):
            add_paragraph(tag.get_text(strip=True))

        processed_tags.append(tag)


def extract_links(soup, base_url):
    """Extract all internal links from page, including navbar links explicitly."""
    links = set()
    # The original soup object is needed here before navbars are decomposed
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        full_url = urljoin(base_url, href)
        parsed = urlparse(full_url)
        # Ensure we only crawl links on the same domain
        if parsed.netloc == urlparse(BASE_URL).netloc:
            # Clean fragments and query parameters
            clean_url = parsed.scheme + "://" + parsed.netloc + parsed.path
            links.add(clean_url)

    return links

def crawl(url, depth=0, max_depth=200):
    if url in visited or depth > max_depth:
        return
    visited.add(url)
    print(f"Crawling: {url} at depth {depth}")
    try:
        driver.get(url)
        time.sleep(WAIT_TIME)
        soup = BeautifulSoup(driver.page_source, "html.parser")
        
        # Extract links first from the original soup
        links = extract_links(soup, url)

        # Now parse and save the content (which modifies the soup by removing navs)
        add_heading(f"Page: {url}", level=1)
        parse_and_save_content(soup)

        for link in links:
            if link not in visited:
                crawl(link, depth + 1, max_depth)
    except Exception as e:
        print(f"Failed URL: {url} due to {e}")

def main():
    crawl(BASE_URL)
    doc.save(WORD_FILENAME)
    driver.quit()
    print(f"Website saved to {WORD_FILENAME}")

if __name__ == "__main__":
    main()
